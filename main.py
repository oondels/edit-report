import tkinter as tk
import ttkbootstrap as ttk
import os
import re

from tkinter import filedialog
from tkinter import messagebox
from ttkbootstrap.constants import *
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches
from datetime import date

filenames = []


def success_popup():
    messagebox.showinfo(
        "Sucesso!", "O relatório foi salvo com sucesso.\nVerifique na pasta documentos!"
    )


def show_error():
    messagebox.showerror("Erro", "Preencha todos os campos antes de prosseguir!")


def editDoc(doc, padrao, novo):
    for paragrafo in doc.paragraphs:
        if padrao in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(padrao, novo)


def addDefects(relatorio, array):
    defects_table = relatorio.tables[0]
    for defect in array:
        newDefect = defects_table.cell(0, 0)
        newDefect.add_paragraph().add_run().add_text(defect)


def addServices(relatorio, array):
    services_table = relatorio.tables[1]
    for service in array:
        newService = services_table.cell(0, 0)
        newService.add_paragraph().add_run().add_text(service)


def addMeasures(relatorio, array):
    measures_table = relatorio.tables[2]

    for measure in array:
        new_meausre = measures_table.cell(0, 0)
        new_meausre.add_paragraph().add_run().add_text(measure)


def selectImages():
    global filenames
    filenames = filedialog.askopenfilenames(
        title="Selecione as imagens",
        filetypes=[("Image files", "*.jpg *.png *.jpeg *.bmp")],
    )

    return filenames


def sendImages(relatorio):
    pictures_table = relatorio.tables[3]

    for i in range(0, len(filenames), 2):
        sub_array = filenames[i : i + 2]
        row = pictures_table.add_row()

        for j in range(len(sub_array)):
            cell = row.cells[j]
            cell.add_paragraph().add_run().add_picture(sub_array[j], width=Inches(5.0))


def saveRelatorio():
    # Load document
    relatorio = Document("./data/relatorio-padrao.docx")

    # Get current date
    data_atual = date.today()

    # Get the entries
    sendImages(relatorio)
    newEquipament = equipament.get().strip()
    newOsNumber = osNumber.get().strip()
    newEmpresa = empresaContratante.get().strip()
    newdate_entrada = date_entrada.entry.get().strip()
    newExecution = execution.entry.get().strip()

    # Split the entries into array
    array_defects = re.split(r"[,\n]", defects.get("1.0", END).strip())
    array_services = re.split(r"[,\n]", services.get("1.0", END).strip())
    array_measures = re.split(r"[,\n]", measure.get("1.0", END).strip())

    # Check if has empty entries
    if (
        (not newEmpresa)
        or (not newEquipament)
        or (not newOsNumber)
        or (not array_defects or (len(array_defects) == 1 and not array_defects[0]))
        or (not array_services or (len(array_services) == 1 and not array_services[0]))
        or (not array_measures or (len(array_measures) == 1 and not array_measures[0]))
        or (not newdate_entrada)
        or (not newExecution)
    ):

        return show_error()

    # Insert the entries
    editDoc(relatorio, "nome-equipamento", newEquipament)
    editDoc(relatorio, "numero-os", newOsNumber)
    editDoc(relatorio, "data-entrada", newdate_entrada)
    editDoc(relatorio, "data-execucao", newExecution)
    editDoc(relatorio, "empresa-contratante", newEmpresa)

    addDefects(relatorio, array_defects)
    addServices(relatorio, array_services)
    addMeasures(relatorio, array_measures)

    success_popup()

    documents_path = os.path.join(os.environ["USERPROFILE"], "Documents")
    path = documents_path + f"\\Relatorios-Editados\\{data_atual}\\"
    filename = path + f"Relatorio-{newEquipament}-{data_atual}.docx"

    # Reset the entries
    equipament.delete(0, END)
    osNumber.delete(0, END)
    empresaContratante.delete(0, END)
    defects.delete("1.0", END)
    services.delete("1.0", END)
    measure.delete("1.0", END)

    if not os.path.exists(path):
        os.makedirs(path)

    relatorio.save(filename)


if __name__ == "__main__":
    app = ttk.Window(themename="superhero", minsize=(800, 850))

    app.title("Relatórios Engenharia")
    app.geometry("800x850")

    # Frame principal
    main_frame = ttk.Frame(app, padding="20")
    main_frame.pack(expand=True)

    # Title Frame
    info_frame = ttk.Frame(main_frame, padding="10", relief="raised", borderwidth=1)
    info_frame.grid(row=0, column=0, columnspan=2, sticky="nsew")

    logo = Image.open("./data/jo-diesel.jpg")
    logo = logo.resize((100, 100))  # Ajuste o tamanho da imagem conforme necessário
    logo_tk = ImageTk.PhotoImage(logo)
    logo_label = ttk.Label(info_frame, image=logo_tk)
    logo_label.image = logo_tk  # Mantenha uma referência da imagem
    logo_label.pack(pady=5)

    labelTitle = ttk.Label(
        info_frame,
        text="Relatórios - Diesel Hidráulica",
        font=("Helvetica", 16, "bold"),
    )
    labelTitle.pack(pady=5)

    engineerName = ttk.Label(
        info_frame,
        text="Engenheiro - Marvin Portela",
        font=("Helvetica", 10, "bold"),
        bootstyle="success",
    )
    engineerName.pack(pady=3)

    # Frame para a primeira coluna
    div_one = ttk.Frame(main_frame, padding="10", relief="raised", borderwidth=1)
    div_one.grid(row=1, column=0, sticky="nsew")

    labelEquipament = ttk.Label(
        div_one, text="Nome do equipamento", font=("Helvetica", 12, "bold")
    )
    labelEquipament.grid(row=0, column=0, padx=5, pady=5, sticky=W)
    equipament = ttk.Entry(div_one, bootstyle="primary", width=30)
    equipament.grid(row=1, column=0, padx=5, pady=5)

    labelOs = ttk.Label(div_one, text="Número de O.S", font=("Helvetica", 12, "bold"))
    labelOs.grid(row=2, column=0, padx=5, pady=5, sticky=W)
    osNumber = ttk.Entry(div_one, bootstyle="primary", width=30)
    osNumber.grid(row=3, column=0, padx=5, pady=5)

    empresaContranteLabel = ttk.Label(
        div_one, text="Empresa Contratante", font=("Helvetica", 12, "bold")
    )
    empresaContranteLabel.grid(row=4, column=0, padx=5, pady=5, sticky=W)
    empresaContratante = ttk.Entry(div_one, bootstyle="primary", width=30)
    empresaContratante.grid(row=5, column=0, padx=5, pady=5)

    dateLabel = ttk.Label(
        div_one, text="Data de entrada", font=("Helvetica", 12, "bold")
    )
    dateLabel.grid(row=6, column=0, padx=5, pady=5, sticky=W)
    date_entrada = ttk.DateEntry(div_one, bootstyle="primary", width=30)
    date_entrada.grid(row=7, column=0, padx=5, pady=5)

    executionLabel = ttk.Label(
        div_one, text="Data de execução", font=("Helvetica", 12, "bold")
    )
    executionLabel.grid(row=8, column=0, padx=5, pady=5, sticky=W)
    execution = ttk.DateEntry(div_one, bootstyle="success", width=30)
    execution.grid(row=9, column=0, padx=5, pady=5)

    # Frame para a segunda coluna
    div_two = ttk.Frame(main_frame, relief=RAISED, borderwidth=1, padding="10")
    div_two.grid(row=1, column=1, sticky="nsew")

    defectsLabel = ttk.Label(div_two, text="Defeitos", font=("Helvetica", 12, "bold"))
    defectsLabel.grid(row=0, column=0, padx=5, pady=5, sticky=W)
    defects = tk.Text(div_two, wrap=tk.WORD, height=3, width=30)
    defects.grid(row=1, column=0, padx=5, pady=5)

    servicesLabel = ttk.Label(div_two, text="Serviços", font=("Helvetica", 12, "bold"))
    servicesLabel.grid(row=2, column=0, padx=5, pady=5, sticky=W)
    services = tk.Text(div_two, wrap=tk.WORD, height=3, width=30)
    services.grid(row=3, column=0, padx=5, pady=5)

    measureLabel = ttk.Label(div_two, text="Medidas", font=("Helvetica", 12, "bold"))
    measureLabel.grid(row=4, column=0, padx=5, pady=5, sticky=W)
    measure = tk.Text(div_two, wrap=tk.WORD, height=3, width=30)
    measure.grid(row=5, column=0, padx=5, pady=5)

    picturesLabel = ttk.Label(div_two, text="Fotos", font=("Helvetica", 12, "bold"))
    picturesLabel.grid(row=6, column=0, padx=5, pady=5, sticky=W)
    pictures = ttk.Button(
        div_two, text="Selecionar fotos", bootstyle="success", command=selectImages
    )
    pictures.grid(row=7, column=0, padx=5, pady=5)

    # Submit Frame
    submitFrame = ttk.Frame(main_frame, padding="10")
    submitFrame.grid(row=2, column=0, columnspan=2, sticky="nsew")

    buttonEnviar = ttk.Button(submitFrame, text="Salvar", command=saveRelatorio)
    buttonEnviar.pack()

    # Ajuste a proporção das colunas e linhas
    app.grid_columnconfigure(0, weight=1)
    app.grid_columnconfigure(1, weight=1)
    app.grid_rowconfigure(1, weight=1)

    app.mainloop()
